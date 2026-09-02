import io
import re
import zipfile

import fitz  # pymupdf
import pytesseract
from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from config import TESSERACT_CMD, GENERATION_MODEL

if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

MIN_PAGE_TEXT_LENGTH = 20


class ExtractionError(Exception):
    status_code = 422
    code = "DOCUMENT_EXTRACTION_FAILED"


class PdfPasswordProtectedError(ExtractionError):
    code = "PDF_PASSWORD_PROTECTED"


class DocumentCorruptedError(ExtractionError):
    code = "DOCUMENT_CORRUPTED"


class OcrUnavailableError(ExtractionError):
    status_code = 503
    code = "OCR_UNAVAILABLE"


class NoReadableContentError(ExtractionError):
    code = "NO_READABLE_CONTENT"


def clean_text(text: str) -> str:
    text = (text or "").replace("\x00", "")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _ocr_image_with_gemini(image: Image.Image) -> str:
    """Vision OCR fallback used only when local Tesseract cannot read the image.

    Keeping this as a fallback preserves fast/offline OCR in normal deployments,
    while scanned PDFs still work on hosts where the Tesseract binary or language
    pack is unavailable (UC08-UI03).
    """
    try:
        from google.genai import types
        from llm_client.llm_client import generate_content, AIProviderError

        buffer = io.BytesIO()
        image.convert("RGB").save(buffer, format="PNG")
        image_part = types.Part.from_bytes(data=buffer.getvalue(), mime_type="image/png")
        response = generate_content(
            [
                "Extract all readable text from this image exactly as written. "
                "Preserve important numbers and identifiers. Return only the transcription.",
                image_part,
            ],
            attempts=3,
        )
        return clean_text(getattr(response, "text", "") or "")
    except Exception as exc:
        raise OcrUnavailableError("OCR engine is unavailable or could not read this image.") from exc


def _ocr_image(image: Image.Image) -> str:
    local_text = ""
    try:
        local_text = clean_text(pytesseract.image_to_string(image, lang="vie+eng"))
    except pytesseract.TesseractNotFoundError:
        return _ocr_image_with_gemini(image)
    except pytesseract.TesseractError as exc:
        # Some deployments have English data installed but not Vietnamese. Keep
        # OCR available in that environment instead of failing the whole service.
        message = str(exc).lower()
        if any(token in message for token in ("language", "traineddata", "tessdata")):
            try:
                local_text = clean_text(pytesseract.image_to_string(image, lang="eng"))
            except (pytesseract.TesseractNotFoundError, pytesseract.TesseractError):
                return _ocr_image_with_gemini(image)
        else:
            return _ocr_image_with_gemini(image)

    # Very short OCR output is usually a scan that Tesseract could not decode.
    # Try Gemini Vision before classifying the source as unreadable.
    if len(local_text) < MIN_PAGE_TEXT_LENGTH:
        try:
            vision_text = _ocr_image_with_gemini(image)
            if len(vision_text) > len(local_text):
                return vision_text
        except OcrUnavailableError:
            pass

    return local_text


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise DocumentCorruptedError("The PDF file is corrupted or unreadable.") from exc

    try:
        if doc.needs_pass:
            raise PdfPasswordProtectedError("Password-protected PDFs are not supported.")

        parts = []
        for page_number, page in enumerate(doc):
            embedded = clean_text(page.get_text())
            if len(embedded) >= MIN_PAGE_TEXT_LENGTH:
                page_text = embedded
            else:
                pixmap = page.get_pixmap(dpi=300)
                image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                page_text = _ocr_image(image)

            if page_text:
                parts.append(f"[Page {page_number + 1}]\n{page_text}")

        result = clean_text("\n\n".join(parts))
        if not result:
            raise NoReadableContentError("No readable text was found in this file.")
        return result
    finally:
        doc.close()


def _iter_docx_blocks(document: _Document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _extract_docx_media(file_bytes: bytes) -> tuple[list[str], bool, bool]:
    """Return (ocr_text_parts, has_media, ocr_unavailable).

    Embedded-image OCR is best-effort for a DOCX that already contains readable
    paragraphs/tables. A missing OCR engine must not make an otherwise readable
    Word document unusable.
    """
    text_parts = []
    has_media = False
    ocr_unavailable = False
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            for name in sorted(archive.namelist()):
                if not name.startswith("word/media/") or name.endswith("/"):
                    continue
                has_media = True
                try:
                    image = Image.open(io.BytesIO(archive.read(name)))
                    ocr_text = _ocr_image(image)
                except UnidentifiedImageError:
                    continue
                except OcrUnavailableError:
                    ocr_unavailable = True
                    continue
                if ocr_text:
                    text_parts.append(f"[Image {name.rsplit('/', 1)[-1]}]\n{ocr_text}")
    except zipfile.BadZipFile as exc:
        raise DocumentCorruptedError("The DOCX file is corrupted or unreadable.") from exc
    return text_parts, has_media, ocr_unavailable


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentCorruptedError("The DOCX file is corrupted or unreadable.") from exc

    parts = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            if text:
                parts.append(text)
        else:
            for row in block.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    parts.append(row_text)

    for section in document.sections:
        header = clean_text("\n".join(p.text for p in section.header.paragraphs))
        footer = clean_text("\n".join(p.text for p in section.footer.paragraphs))
        if header:
            parts.append(f"[Header]\n{header}")
        if footer:
            parts.append(f"[Footer]\n{footer}")

    media_parts, has_media, ocr_unavailable = _extract_docx_media(file_bytes)
    parts.extend(media_parts)

    result = clean_text("\n\n".join(parts))
    if result:
        return result

    # A scan-only DOCX still needs OCR. Only in that case should an unavailable
    # OCR engine fail the whole extraction.
    if has_media and ocr_unavailable:
        raise OcrUnavailableError("OCR is required to read the images in this DOCX, but the OCR engine is unavailable.")

    raise NoReadableContentError("No readable text was found in this file.")


def extract_text_from_image(file_bytes: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(file_bytes))
    except (UnidentifiedImageError, OSError) as exc:
        raise DocumentCorruptedError("The image file is corrupted or unreadable.") from exc
    result = _ocr_image(image)
    if not result:
        raise NoReadableContentError("No readable text was found in this file.")
    return result
